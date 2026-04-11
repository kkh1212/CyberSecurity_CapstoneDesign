"""
문서 파싱 및 청킹
출처: 조원 코드 src/chunking.py 기반
변경: HWP/HWPX 제거, src.config 의존성 제거, source 메타데이터 단순화
"""
import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional

import pdfplumber
from docx import Document
from pypdf import PdfReader

CHUNK_SIZE = 500
CHUNK_OVERLAP = 120

TITLE_FIELDS = {"교과목명", "프로그램명", "사업명"}
POLICY_DOC_KEYWORDS = ("규정", "지침", "세칙", "학칙", "방침", "절차", "기준", "규칙")
POLICY_SPLITTABLE_BLOCK_TYPES = {"page_text", "paragraph", "doc_text", "text_section"}
CLAUSE_LINE_PATTERN = re.compile(
    r"^(?:제\s*\d+\s*(?:장|절|조)(?:\s*\([^)]+\))?|부칙|별표\s*\d*|별지\s*\d*)"
)
INLINE_CLAUSE_PATTERN = re.compile(
    r"(?=(?:제\s*\d+\s*(?:장|절|조)(?:\s*\([^)]+\))?|부칙|별표\s*\d*|별지\s*\d*))"
)
SUBCLAUSE_LINE_PATTERN = re.compile(r"^(?:\(?\d+\)?[.)]?|[가-힣A-Za-z][.)])\s+")


# ─── 텍스트 분할 ───────────────────────────────────────────────

def split_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((chunk, start, end))
        if end == text_len:
            break
        start = max(0, end - overlap)
    return chunks


# ─── 테이블 정규화 ─────────────────────────────────────────────

def normalize_table(rows: List[List[str]], label: str) -> str:
    cleaned_rows = [
        [(cell or "").replace("\n", " ").strip() for cell in row]
        for row in rows
        if any((cell or "").strip() for cell in row)
    ]
    if not cleaned_rows:
        return ""

    header = cleaned_rows[0]
    lines = [label] + [" | ".join(cell or "-" for cell in row) for row in cleaned_rows]

    if len(cleaned_rows) > 1 and any(header):
        lines.append("[TABLE ROWS]")
        for row in cleaned_rows[1:]:
            pairs = []
            for idx, cell in enumerate(row):
                key = header[idx] if idx < len(header) else f"column_{idx + 1}"
                if key and cell:
                    pairs.append(f"{key}: {cell}")
            if pairs:
                lines.append("; ".join(pairs))

    return "\n".join(lines).strip()


def build_table_blocks(rows: List[List[str]], label: str) -> List[Dict]:
    cleaned_rows = [
        [(cell or "").replace("\n", " ").strip() for cell in row]
        for row in rows
        if any((cell or "").strip() for cell in row)
    ]
    if not cleaned_rows:
        return []

    blocks: List[Dict] = []
    table_text = normalize_table(cleaned_rows, label)
    if table_text:
        blocks.append({"block_type": "table", "text": table_text})

    header = cleaned_rows[0]
    if len(cleaned_rows) <= 1 or not any(header):
        return blocks

    for row_index, row in enumerate(cleaned_rows[1:], start=1):
        pairs = []
        entity_title = None
        for idx, cell in enumerate(row):
            key = header[idx] if idx < len(header) else f"column_{idx + 1}"
            if not key or not cell:
                continue
            pairs.append(f"{key}: {cell}")
            if key.strip() in TITLE_FIELDS:
                entity_title = cell
        if not pairs:
            continue

        row_text = "\n".join([
            f"[ROW {row_index}]",
            " | ".join(cell or "-" for cell in row),
            "; ".join(pairs),
        ]).strip()
        blocks.append({
            "block_type": "table_row",
            "row_index": row_index,
            "entity_title": entity_title,
            "text": row_text,
        })

    return blocks


# ─── 소스 메타데이터 ───────────────────────────────────────────

def build_source_metadata(path: Path) -> Dict[str, str]:
    return {
        "source": path.name,
        "file_name": path.name,
        "domain": "company",
    }


# ─── 블록 정규화 ────────────────────────────────────────────────

def normalize_block_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_entity_title(text: str) -> Optional[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:3]:
        clean = line.lstrip("-").strip()
        if ":" in clean:
            key, value = clean.split(":", 1)
            if key.strip() in TITLE_FIELDS and value.strip():
                return value.strip()
        match = re.match(r"^\d+\.\s*(.+)$", clean)
        if match:
            return match.group(1).strip()
    if lines and len(lines[0]) <= 40 and ":" not in lines[0]:
        return lines[0]
    return None


# ─── 정책 문서 분할 ────────────────────────────────────────────

def is_policy_document(path: Path, text: str = "") -> bool:
    stem = path.stem
    if any(keyword in stem for keyword in POLICY_DOC_KEYWORDS):
        return True
    normalized = normalize_block_text(text)
    if not normalized:
        return False
    if CLAUSE_LINE_PATTERN.search(normalized):
        return True
    return bool(re.search(r"제\s*\d+\s*조", normalized))


def extract_clause_title(text: str) -> Optional[str]:
    for line in normalize_block_text(text).splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if CLAUSE_LINE_PATTERN.match(stripped):
            return stripped[:120]
        break
    return None


def split_long_policy_section(text: str) -> List[str]:
    normalized = normalize_block_text(text)
    if len(normalized) <= CHUNK_SIZE * 2:
        return [normalized]

    lines = normalized.splitlines()
    sections: List[List[str]] = []
    current: List[str] = []

    def flush():
        nonlocal current
        if current:
            sections.append(current)
            current = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current and len("\n".join(current)) >= CHUNK_SIZE:
                flush()
            elif current:
                current.append("")
            continue
        if current and SUBCLAUSE_LINE_PATTERN.match(stripped) and len("\n".join(current)) >= CHUNK_SIZE:
            flush()
        current.append(stripped)
    flush()

    flattened = ["\n".join(s).strip() for s in sections if any(p.strip() for p in s)]
    return flattened if len(flattened) > 1 else [normalized]


def split_policy_sections(text: str) -> List[str]:
    normalized = normalize_block_text(text)
    if not normalized:
        return []

    prepared = normalized
    if "\n" not in prepared:
        prepared = INLINE_CLAUSE_PATTERN.sub("\n", prepared).strip()

    lines = prepared.splitlines()
    sections: List[List[str]] = []
    current: List[str] = []

    def flush():
        nonlocal current
        if current:
            sections.append(current)
            current = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                current.append("")
            continue
        if CLAUSE_LINE_PATTERN.match(stripped):
            flush()
        current.append(stripped)
    flush()

    flattened = ["\n".join(s).strip() for s in sections if any(p.strip() for p in s)]
    if len(flattened) <= 1:
        fallback = [s.strip() for s in re.split(r"\n\s*\n", normalized) if s.strip()]
        flattened = fallback or [normalized]

    result: List[str] = []
    for section in flattened:
        result.extend(split_long_policy_section(section))
    return [s for s in result if s]


def split_block_for_chunking(block: Dict, path: Path) -> List[Dict]:
    text = (block.get("text") or "").strip()
    if not text:
        return []
    block_type = block.get("block_type", "text")
    if block_type not in POLICY_SPLITTABLE_BLOCK_TYPES or not is_policy_document(path, text):
        return [block]
    sections = split_policy_sections(text)
    if len(sections) <= 1:
        return [block]
    result = []
    for idx, section in enumerate(sections, start=1):
        updated = dict(block)
        updated["text"] = section
        updated["block_type"] = "clause_section"
        updated["sub_block_index"] = idx
        updated["clause_title"] = extract_clause_title(section)
        result.append(updated)
    return result


# ─── 청킹 ─────────────────────────────────────────────────────

def chunk_blocks(blocks: Iterable[Dict], path: Path) -> List[Dict]:
    all_chunks = []
    chunk_index = 0
    source_meta = build_source_metadata(path)

    for block_index, block in enumerate(blocks):
        for sub_block in split_block_for_chunking(block, path):
            text = (sub_block.get("text") or "").strip()
            if not text:
                continue

            block_type = sub_block.get("block_type", "text")
            prefix = f"[SOURCE={source_meta['source']}] [BLOCK={block_type}]"
            if "page" in sub_block:
                prefix += f" [PAGE={sub_block['page']}]"
            if "table_index" in sub_block:
                prefix += f" [TABLE={sub_block['table_index']}]"
            if sub_block.get("clause_title"):
                prefix += f" [CLAUSE={sub_block['clause_title']}]"

            block_text = f"{prefix}\n{text}"

            chunk_base = {}
            for key in ("page", "table_index", "row_index", "block_num",
                        "entity_title", "sub_block_index", "clause_title"):
                if key in sub_block and sub_block[key] is not None:
                    chunk_base[key] = sub_block[key]

            if block_type in {"table", "table_row", "text_section", "clause_section"} \
                    and len(block_text) <= CHUNK_SIZE * 2:
                all_chunks.append({
                    "chunk_id": f"{source_meta['source']}::chunk_{chunk_index}",
                    "source": source_meta["source"],
                    "file_name": source_meta["file_name"],
                    "domain": source_meta["domain"],
                    "file_type": path.suffix.lower(),
                    "text": block_text,
                    "start": 0,
                    "end": len(block_text),
                    "block_type": block_type,
                    "block_index": block_index,
                    **chunk_base,
                })
                chunk_index += 1
                continue

            for _, start, end in split_text(block_text):
                chunk_text = block_text[start:end].strip()
                if not chunk_text:
                    continue
                all_chunks.append({
                    "chunk_id": f"{source_meta['source']}::chunk_{chunk_index}",
                    "source": source_meta["source"],
                    "file_name": source_meta["file_name"],
                    "domain": source_meta["domain"],
                    "file_type": path.suffix.lower(),
                    "text": chunk_text,
                    "start": start,
                    "end": end,
                    "block_type": block_type,
                    "block_index": block_index,
                    **chunk_base,
                })
                chunk_index += 1

    return all_chunks


# ─── 파일 로더 ─────────────────────────────────────────────────

def load_txt_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp949"):
        try:
            return path.read_text(encoding=encoding).strip()
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def load_txt_blocks(path: Path) -> List[Dict]:
    text = load_txt_file(path)
    if not text:
        return []
    sections = [s.strip() for s in re.split(r"\n\s*\n", text) if s.strip()]
    if len(sections) <= 1:
        return [{"block_type": "text_section", "text": text,
                 "entity_title": extract_entity_title(text)}]
    return [
        {"block_type": "text_section", "block_num": idx, "text": section,
         "entity_title": extract_entity_title(section)}
        for idx, section in enumerate(sections, start=1)
    ]


def extract_pdf_blocks(path: Path) -> List[Dict]:
    blocks = []
    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append({"block_type": "page_text", "page": page_num, "text": text})
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for table_idx, table in enumerate(tables, start=1):
                for block in build_table_blocks(table, f"[TABLE {table_idx}]"):
                    block["page"] = page_num
                    block["table_index"] = table_idx
                    blocks.append(block)
    return blocks


def extract_pdf_with_pypdf(path: Path) -> List[Dict]:
    reader = PdfReader(str(path))
    blocks = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            page_text = (page.extract_text() or "").strip()
        except Exception:
            page_text = ""
        if page_text:
            blocks.append({"block_type": "page_text", "page": page_num, "text": page_text})
    return blocks


def load_pdf_blocks(path: Path) -> List[Dict]:
    try:
        blocks = extract_pdf_blocks(path)
        if blocks:
            return blocks
    except Exception:
        pass
    return extract_pdf_with_pypdf(path)


def load_docx_blocks(path: Path) -> List[Dict]:
    document = Document(str(path))
    blocks = []
    for idx, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            blocks.append({"block_type": "paragraph", "block_num": idx, "text": text})
    for table_idx, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        for block in build_table_blocks(rows, f"[TABLE {table_idx}]"):
            block["table_index"] = table_idx
            blocks.append(block)
    return blocks


def load_document_blocks(path: Path) -> List[Dict]:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return load_txt_blocks(path)
    if suffix == ".pdf":
        return load_pdf_blocks(path)
    if suffix == ".docx":
        return load_docx_blocks(path)
    return []


def load_documents_from_dirs(doc_dirs: List[Path]) -> List[Dict]:
    """여러 폴더에서 문서를 로드해 청크 목록 반환"""
    supported = {".txt", ".pdf", ".docx"}
    excluded_tokens = ("question", "questions", "질문예시", "예시질문")
    all_chunks = []

    for docs_dir in doc_dirs:
        if not docs_dir.exists():
            continue
        file_paths = sorted(
            fp for fp in docs_dir.rglob("*")
            if fp.is_file()
            and fp.suffix.lower() in supported
            and not any(tok in fp.stem.lower() for tok in excluded_tokens)
        )
        for fp in file_paths:
            blocks = load_document_blocks(fp)
            all_chunks.extend(chunk_blocks(blocks, fp))

    return all_chunks
