"""
Convert experiment attack .txt files to .docx while preserving paragraph breaks.

This is for format-sensitivity experiments: same payload text, different document
container. The output directory mirrors the source attack corpus layout.

Usage:
  python -m experiments.attack.convert_txt_to_docx
  python -m experiments.attack.convert_txt_to_docx \
      --src data/exp_corpus/attack --dst data/exp_corpus/attack_docx
"""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SRC = PROJECT_ROOT / "data" / "exp_corpus" / "attack"
DEFAULT_DST = PROJECT_ROOT / "data" / "exp_corpus" / "attack_docx"


def write_docx_from_txt(src: Path, dst: Path) -> None:
    text = src.read_text(encoding="utf-8")
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10.5)

    for block in text.split("\n\n"):
        paragraph_text = block.strip()
        if not paragraph_text:
            continue
        paragraph = document.add_paragraph()
        for idx, line in enumerate(paragraph_text.splitlines()):
            if idx:
                paragraph.add_run().add_break()
            paragraph.add_run(line)

    dst.parent.mkdir(parents=True, exist_ok=True)
    document.save(dst)


def convert_corpus(src_root: Path, dst_root: Path) -> int:
    if not src_root.exists():
        raise FileNotFoundError(f"Source attack corpus not found: {src_root}")
    if dst_root.exists():
        shutil.rmtree(dst_root)
    dst_root.mkdir(parents=True, exist_ok=True)

    count = 0
    for src in sorted(src_root.rglob("*.txt")):
        rel = src.relative_to(src_root)
        dst = dst_root / rel.with_suffix(".docx")
        write_docx_from_txt(src, dst)
        count += 1
    return count


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert attack txt corpus to docx.")
    parser.add_argument("--src", type=Path, default=DEFAULT_SRC)
    parser.add_argument("--dst", type=Path, default=DEFAULT_DST)
    args = parser.parse_args()

    count = convert_corpus(args.src, args.dst)
    print(f"[convert_txt_to_docx] converted {count} files")
    print(f"[convert_txt_to_docx] dst={args.dst}")


if __name__ == "__main__":
    main()
